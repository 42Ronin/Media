#!/usr/bin/env python3
"""Builds the Lesson 1 script — Why HMIS — in the format the team established.

    python3 script/make_script.py   -> Script - HMIS BNTS Lesson 1 - Why HMIS.docx

FIRST DRAFT. Nothing in it has been through a sourcing pass: the research that
fed it was gathered through search summaries because direct page fetches are
blocked from the build environment, so no primary document was read end to end.
Every factual claim carries a VERIFY note naming what has to be checked and
against what. Read those before this goes to anyone outside the project.

Conventions follow NHO.docx: bold+underline sections, italic production notes,
red for questions back to the author. Slide numbers are gone, per the revision
that dropped them from the Search script.
"""
import os

import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(HERE, "..", "Script - HMIS BNTS Lesson 1 - Why HMIS.docx")

GREEN = RGBColor(0x1E, 0x7A, 0x33)
ORANGE = RGBColor(0xC0, 0x60, 0x00)
BLUE = RGBColor(0x1F, 0x4E, 0xD8)
RED = RGBColor(0xC0, 0x1C, 0x1C)
GREY = RGBColor(0x5F, 0x77, 0x87)

d = docx.Document()
st = d.styles["Normal"]
st.font.name = "Calibri"
st.font.size = Pt(11)
for s in d.sections:
    s.left_margin = s.right_margin = Inches(1)


def para(text="", bold=False, italic=False, underline=False, color=None,
         size=None, space_after=6):
    p = d.add_paragraph()
    if text:
        r = p.add_run(text)
        r.bold, r.italic, r.underline = bold, italic, underline
        if color is not None:
            r.font.color.rgb = color
        if size is not None:
            r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def title(t):
    return para(t, bold=True, underline=True, size=15, space_after=12)


def section(t):
    d.add_page_break()
    return para(t, bold=True, underline=True, size=14, space_after=10)


def slide(t):
    return para(t, bold=True, size=12.5, space_after=6)


def lab(t):
    return para(t, bold=True, size=9.5, color=GREY, space_after=2)


def body(t):
    return para(t, space_after=8)


def note(t):
    return para(t, italic=True, space_after=8)


def ask(t):
    return para(t, color=RED, space_after=8)


def verify(t):
    return para("VERIFY — " + t, color=RED, italic=True, size=10, space_after=8)


def bullet(t):
    p = d.add_paragraph(style="List Bullet")
    p.add_run(t)
    p.paragraph_format.space_after = Pt(3)
    return p


def onscreen(lines):
    lab("ON SCREEN")
    for l in lines:
        bullet(l)
    para(space_after=6)


# ══════════════════════════════════════════════════════════ front matter
title("HMIS BNTS")
para("Lesson 1: Why HMIS — Elearning Course Script", bold=True, size=13, space_after=4)
para("First draft, for internal review. Not sourced.", italic=True, color=GREY, space_after=14)

para("About This Document", bold=True, underline=True, space_after=6)
body("This is the first lesson of the series. Search has moved to Lesson 2. This one is a "
     "prologue: it answers the question a new user actually arrives with, which is not how to "
     "use the software but why they are being asked to. There is no narrative in it — the "
     "series' story begins in Lesson 2, with Desmond.")
body("It is written to be read aloud. Narration is the spoken text. On-screen text is what the "
     "learner reads. Production notes are in italics and describe what a screen does.")

para("What Is Being Asked of the Reviewer", bold=True, underline=True, space_after=6)
body("Three things.")
bullet("Whether the argument is the right one — that HMIS is worth caring about because of what "
       "it does to a person's chances, not because it is required.")
bullet("Whether the tone is right for someone who has been handed a login and told to fill it in.")
bullet("Every red line. Those are either questions for you or claims that need a source before "
       "this goes any further.")
para(space_after=8)

para("Sourcing — read this before the content", bold=True, underline=True, color=RED, space_after=6)
body("Nothing here has been verified against a primary document. The research behind it was "
     "gathered from search summaries, because direct page retrieval is blocked from the "
     "environment this was written in — hudexchange.info, lahsa.org, rand.org and arxiv.org all "
     "refused the fetch.")
body("That means every figure below is second-hand. The numbers are probably right; I have not "
     "read one of them in its own document. Anything marked VERIFY needs someone with working "
     "access to open the source and confirm the wording before this is shown outside the team.")
ask("Who owns the sourcing pass, and do they have access to LAHSA's Data Quality Monitoring "
    "Plan? Two of the strongest lines in this lesson depend on how LAHSA itself frames the "
    "stakes, and that document could not be opened from here.")

# ══════════════════════════════════════════════════════════ objectives
section("Course Overview")

slide("Objectives")
body("By the end of this lesson, you will be able to:")
bullet("Say what HMIS is and who else can see what you put in it.")
bullet("Describe what a participant's record does for them.")
bullet("Explain how the information you enter reaches a decision about housing and funding.")
bullet("Recognise what a duplicate record costs the person it belongs to.")
bullet("State what a participant can refuse, and what you must never make conditional on their "
       "consent.")

slide("Lesson Structure")
onscreen([
    "Why this lesson comes first",
    "What you are typing into",
    "What the record does",
    "Where the information goes",
    "What breaks",
    "What the system will not do",
    "What this means tomorrow",
])
note("Roughly twelve to fifteen minutes narrated. No simulation in this lesson — the software "
     "arrives in Lesson 2, and meeting it here would answer a question nobody has asked yet.")

# ══════════════════════════════════════════════════════════ prologue
section("Before Any of It")

slide("Why This Lesson Comes First")
lab("NARRATION")
body("You will be given a login and asked to enter information about people. Most people are "
     "shown the buttons and left to work out the point of it later, or never.")
body("This lesson is the point of it. It is short, there is no software in it, and nothing here "
     "asks you to do anything yet. It is here so that when the rest of the series shows you how, "
     "you already know why it is worth the care.")
onscreen(["No software in this lesson. Just what it is for."])
note("Prologue, deliberately plain. The series' story starts in Lesson 2 and should not be "
     "borrowed here — no scene, no participant, nobody asking for anything.")

# ══════════════════════════════════════════════════════════ what it is
section("What You Are Typing Into")

slide("One Sentence")
lab("NARRATION")
body("HMIS stands for Homeless Management Information System. It is the shared record of who "
     "has been helped, by whom, and what happened.")
body("That is the whole definition. The rest of this lesson is what follows from the word "
     "shared.")
onscreen(["HMIS — the shared record of who has been helped, by whom, and what happened."])
note("Hold this on screen. It is the only definition the lesson gives, and everything after it "
     "is consequence rather than vocabulary.")

slide("Shared Means Shared")
lab("NARRATION")
body("Your agency is not the only one looking. Los Angeles runs one HMIS across the "
     "Continuum of Care, and the outreach worker who logged that contact months ago may have "
     "worked for an organisation you have never heard of.")
body("This cuts both ways, and both ways matter. What you write can be read by the next worker "
     "who meets this person. What they wrote is sitting there for you, if you can find it.")
verify("LAHSA is the lead agency for the LA Continuum of Care and operates the HMIS. The "
       "implementation is shared across four CoCs — LAHSA, Pasadena, Glendale and Orange "
       "County — as the LA/OC HMIS Collaborative. Confirm against LAHSA's own description "
       "before stating the four-CoC arrangement in narration.")
verify("The software is Clarity Human Services, live for LAHSA since 1 June 2017. Confirm the "
       "date from LAHSA's own announcement.")
ask("Do we say how many agencies and users share the system? It is the single most persuasive "
    "number for 'shared means shared', and I could not find a figure I trust. If LAHSA has a "
    "current count, it belongs here.")

# ══════════════════════════════════════════════════════════ what the record does
section("What the Record Does")

slide("Four Jobs")
lab("NARRATION")
body("A record is not a form you filled in. It does four things for the person it belongs to, "
     "and it does them whether or not anyone thinks about it.")
onscreen([
    "It carries their history, so they do not have to keep retelling it.",
    "It can prove what they have been through, when they have nothing left to prove it with.",
    "It decides what they are offered, and when.",
    "It counts.",
])

slide("It Carries Their History")
lab("NARRATION")
body("Ask someone to account for the last two years of their life, from memory, standing on a "
     "pavement. Then ask them again at the next agency. Then again when their case manager "
     "leaves.")
body("A record that can be found is the difference between being known and starting over. That "
     "is not an administrative convenience. Being asked to relive the worst part of your life "
     "for a stranger with a clipboard is a cost, and a findable record is how you stop charging "
     "it to the same person over and over.")
verify("This argument is standard practitioner reasoning and appears across CoC training "
       "materials, but no published study was found that quantifies the burden of repeated "
       "intake. Either soften to experience-based language or find a citation. Do not present "
       "it as a research finding.")
note("Trauma-informed framing matters here. The line is that the record spares the retelling, "
     "not that the worker should extract more detail to make the record fuller.")

slide("It Can Prove What They Have Been Through")
lab("NARRATION")
body("Documents get lost, stolen, soaked, or left behind in a sweep. A record can stand in "
     "their place. A shelter stay you logged, or a contact an outreach worker recorded, can "
     "serve as third-party documentation of somebody's homelessness — which is what a housing "
     "programme asks for before it will let them in.")
body("Somebody typed that contact in a hurry, at the side of a road, eighteen months ago. It is "
     "now the evidence.")
verify("The vetted draft (script_l1_v3, slide 1.8) makes this claim and marks it 'Sourced; a "
       "source link goes here in the sourcing pass' — meaning it was accepted internally but "
       "never cited. This is the strongest single argument in the lesson and it is the one with "
       "no source attached. It needs one.")

slide("It Decides What They Are Offered")
lab("NARRATION")
body("When a housing unit becomes available, nobody walks the streets looking for a suitable "
     "person. They work from a list. That list is built from what has been entered — who is "
     "known, what has been assessed, what they are eligible for.")
body("If someone is not on it, or is on it twice, or is on it under a name nobody can match to "
     "the person standing in front of you, the unit goes to somebody else. Not out of malice. "
     "The list is what there is.")
verify("Coordinated Entry practice varies by CoC, and prioritisation methods have changed — "
       "the VI-SPDAT is increasingly used as an eligibility screen rather than a strict ranked "
       "queue, following published findings of racial bias in the tool. Confirm how LAHSA's CE "
       "works NOW before describing a mechanism. Describe it as 'a list built from what is "
       "entered' unless someone confirms the specifics.")

slide("It Counts")
lab("NARRATION")
body("Once a year, this region counts how many people are homeless. That number is not "
     "trivia. It is the number that decides where money goes and how much of it there is.")
body("Your records are part of how that count is made.")
verify("HMIS is a direct data source for the sheltered Point-in-Time count and the Housing "
       "Inventory Count. Confirm against HUD's PIT/HIC guidance.")

# ══════════════════════════════════════════════════════════ the chain
section("Where the Information Goes")

slide("The Chain, Once")
lab("NARRATION")
body("Here is the whole route, once, and then we will not do this again.")
body("You enter a record. The record puts a person on the list that housing is offered from. "
     "The same records are counted, once a year, to say how many people this region is "
     "responsible for. Those counts and the outcomes attached to them become the measures the "
     "region is judged on. And those measures are part of what decides how much federal money "
     "comes back next year — to this region, and to programmes like the one you work for.")
body("Nobody expects you to hold that chain in your head while you work. You only need the "
     "first link. "
     "The record is where a person enters the system that is supposed to help them.")
onscreen([
    "Record → the list housing is offered from",
    "→ the annual count",
    "→ the measures the region is judged on",
    "→ next year's funding",
])
verify("Each link was found separately and no single document states the chain end to end — it "
       "is a synthesis. Every link needs its own confirmation, and the joins are where an "
       "overstatement would creep in. In particular, confirm that data quality itself is scored "
       "in the annual funding competition, which is what makes the last link real rather than "
       "rhetorical.")
note("Build the chain one link at a time on click. It is the only diagram in the lesson and it "
     "is doing the work of about four slides.")
ask("Is this too much? The alternative is to stop after the second link — the list and the "
    "count — and leave funding out. Funding is true and it is what management cares about, but "
    "it is also the part that makes a worker feel like a data-entry clerk. My instinct is to "
    "keep it and keep it short. Yours?")

slide("What the Count Decides")
lab("NARRATION")
body("It is worth knowing how wrong a count can be.")
body("An independent research team counted the same streets as the official count and found the "
     "official figure was missing a large share of the people who were actually there. In Skid "
     "Row, the official count found around six in ten. Extrapolated across the city, thousands "
     "of people may be missing from the number that decides where help is sent.")
body("Their warning was blunt: an inaccurate count can divert resources away from the "
     "communities that need them most.")
onscreen(["Skid Row, 2025: the official count found roughly 6 in 10 of the people an "
          "independent count located."])
verify("RAND Corporation, 2025 (report RRA4438-1). Skid Row 61%, Hollywood 81%, Venice 76% of "
       "RAND's independent LA LEADS count; up to ~7,900 people and dwellings possibly missing "
       "citywide. Covered by LAist and CBS LA. CONFIRM ALL FIGURES — none was read in the "
       "original report.")
ask("Important honesty check. RAND's finding is about the physical street count missing people, "
    "NOT about duplicate records in HMIS. It is the best illustration I found of 'bad numbers "
    "move resources', but if we use it we have to be careful not to imply it is about data "
    "entry. The narration above says 'a count can be wrong' rather than 'your typing caused "
    "this'. Are you comfortable with that line, or would you rather cut it and lose the "
    "strongest concrete example in the lesson?")

# ══════════════════════════════════════════════════════════ what breaks
section("What Breaks")

slide("One Person, Two Records")
lab("NARRATION")
body("Here is the failure this whole series exists to prevent.")
body("Somebody searches for a participant. They do not find them — a nickname instead of a legal "
     "name, a spelling nobody could have guessed, a missing middle initial. So they create a new "
     "record, because that is the helpful thing to do when someone is standing in front of you "
     "waiting.")
body("Now there are two. And everything the first one held is invisible to whoever is looking at "
     "the second. Their existing consent form. Their case manager, who they already trust. The "
     "denial they already appealed. The time they have already spent waiting.")
body("They have not lost their place in the queue. They have lost the queue.")
onscreen(["A duplicate does not split someone's history in half.",
          "It hides one half from whoever is looking at the other."])

slide("And It Does Not Fix Itself")
lab("NARRATION")
body("Nothing stops you creating a second record. No warning appears. The system will let it "
     "happen and carry on, and the person who finds out is usually the participant, weeks later, "
     "when something they were counting on does not arrive.")
body("Untangling it is somebody's job for weeks. They wait through all of it.")
verify("Two sources conflict. Help-desk guidance in the HMIS vendor's own network says "
       "duplicates do not block data entry and are not subject to a HUD de-duplication mandate. "
       "Against that, the statutory purpose of HMIS is an unduplicated count, and data quality "
       "is scored. The narration above avoids the compliance question entirely and stays on what "
       "a duplicate does to the person — which I believe is both true and the better argument. "
       "Confirm with LAHSA whether we can say anything stronger.")
note("This slide is the seam. Lesson 2 is the whole answer to it: search properly, and do not "
     "create until you have.")

# ══════════════════════════════════════════════════════════ the honest part
section("What the System Will Not Do")

slide("Three Honest Things")
lab("NARRATION")
body("If this lesson only told you the system was important, you would be right not to trust the "
     "rest of it. So, three things it does not do.")
body("A participant can say no. They can decline to have their information entered or shared, "
     "and you cannot make services conditional on their agreeing. Ever. Entering someone's data "
     "is not the price of being helped.")
body("Some people are kept out of it on purpose. Survivors of domestic violence are not entered "
     "into the shared system at all — victim service providers use a separate database, by law, "
     "for their safety. The system is deliberately built with people missing from it.")
body("And a clean record does not conjure a housing unit that does not exist. Good data does not "
     "make the shortage go away. It decides who is visible when something does become "
     "available — which matters enormously, and is not the same as a promise.")
onscreen([
    "They can refuse. Services are never conditional on consent.",
    "Survivors are kept out of the shared system, by law.",
    "Good data does not create housing. It decides who is visible when there is some.",
])
verify("Consent, the right to refuse without losing services, and the right to revoke are "
       "documented across CoC policy manuals and LAHSA's knowledge base is reported to hold a "
       "Revocation of Consent form. The DV carve-out is VAWA-based: victim service providers "
       "must use a comparable database rather than HMIS. Confirm all three against LAHSA's own "
       "policy before this is narrated — this is the section where being wrong would be worst.")
ask("Is this section a risk? It tells a new worker, in their first lesson, that participants can "
    "refuse. That is true, and saying it plainly is what earns their trust for everything else. "
    "But it may not be what a programme manager expects lesson one to say. I would keep it. "
    "Confirm before it goes to a wider review.")

# ══════════════════════════════════════════════════════════ close
section("What This Means Tomorrow")

slide("Four Habits")
lab("NARRATION")
body("Nothing in this lesson asks you to do anything differently yet. It asks you to expect "
     "something: that the person in front of you probably already exists in there.")
body("The rest of the series is the how. It comes down to four habits.")
onscreen([
    "Assume a record already exists.",
    "Search before you create — properly, more than once, more than one way.",
    "Ask what else they have been called.",
    "When you cannot find them, that is not proof they are new.",
])

slide("Lesson Closing")
lab("NARRATION")
body("That is the whole argument. What you type is the only version of a person this system "
     "holds, and it decides what they are offered, what they can prove, and whether anyone can "
     "find them again.")
body("The next lesson is where the work starts: finding the record that is already there.")
onscreen(["Next: finding a participant's record."])

# ══════════════════════════════════════════════════════════ kc
section("Knowledge Check")

note("Six questions, three answers each, one correct. Same card format as Lesson 2. Draft only — "
     "these follow the argument above and should be rewritten once the sourcing pass settles "
     "what the lesson is allowed to claim.")

QS = [
    ("A participant tells you nobody has ever taken their details. What should you assume?",
     [("That a record may well exist anyway, and it is worth searching for", True),
      ("That they are new to the system and you should create a record", False),
      ("That they are mistaken and you should press them on it", False)],
     "People forget, and they are often recorded by an agency they never knew the name of. "
     "What they remember is not evidence of what is in the system."),
    ("Who can see what you enter?",
     [("Any agency in the Continuum of Care that shares the system", True),
      ("Only your own agency", False),
      ("Only your supervisor and HMIS support", False)],
     "It is one shared record. That is what makes it worth having, and why what you type "
     "outlives your involvement."),
    ("A participant refuses to have their information entered. What happens to their services?",
     [("Nothing. Services are never conditional on consent", True),
      ("They cannot be enrolled until they agree", False),
      ("They can be served once, but not again", False)],
     "They can refuse, and they can change their mind later. Being entered into the system is "
     "not the price of being helped."),
    ("What does a duplicate record cost the participant?",
     [("Their history becomes invisible to whoever is looking at the other record", True),
      ("Nothing — the system merges them automatically", False),
      ("A short delay while someone tidies it up", False)],
     "Their consent form, their case manager, their waiting time and their prior decisions are "
     "all sitting in a record the person in front of them cannot see."),
    ("Why does the annual count matter to your work?",
     [("It is part of how funding is decided for the whole region", True),
      ("It is published for the media", False),
      ("It sets each agency's individual targets", False)],
     "The count feeds the measures the region is judged on, and those are part of what decides "
     "how much money comes back."),
    ("A record can serve as documentation of somebody's homelessness. Why does that matter?",
     [("Because they may have lost every document they had", True),
      ("Because it saves the agency printing costs", False),
      ("Because HUD requires a paper file", False)],
     "A logged contact or a shelter stay can stand in for paperwork that was lost, stolen or "
     "left behind — and it is often what a housing programme asks for."),
]

for n, (q, answers, fb) in enumerate(QS, 1):
    slide("Question %d" % n)
    body(q)
    for letter, (txt, ok) in zip("ABC", answers):
        para("%s. %s" % (letter, txt), color=GREEN if ok else ORANGE, space_after=3)
    lab("FEEDBACK")
    body(fb)

# ══════════════════════════════════════════════════════════ appendix
section("Appendix — What Still Needs a Source")

body("Ordered by how badly the lesson needs it.")

for claim, why in [
    ("A record can serve as third-party documentation of homelessness.",
     "The strongest argument in the lesson. Carried over from the vetted draft, where it is "
     "marked as sourced but has no link attached."),
    ("RAND's 2025 undercount figures for Los Angeles.",
     "The only concrete example in the lesson. If it cannot be confirmed, the slide goes."),
    ("How LAHSA's Coordinated Entry actually prioritises today.",
     "The narration deliberately says 'a list built from what is entered' because the mechanism "
     "has changed and varies. If someone can confirm the current practice, this can be made "
     "specific and much stronger."),
    ("Consent, the right to refuse, revocation, and the DV carve-out.",
     "Being wrong here would be worse than being vague. Needs LAHSA policy, not general guidance."),
    ("That data quality is scored in the annual funding competition.",
     "This is the link that makes the chain end in something real. Without it the funding claim "
     "is hand-waving."),
    ("How many agencies and users share the LA HMIS.",
     "Not found. Would be the most persuasive single number in the lesson if LAHSA publishes it."),
    ("Clarity live date, and the four-CoC collaborative.",
     "Minor, but stated as fact in narration, so it should be right."),
]:
    para(claim, bold=True, space_after=2)
    para(why, italic=True, color=GREY, space_after=8)

d.save(OUT)
print("  %s" % os.path.basename(OUT))
print("  %d paragraphs" % len(d.paragraphs))
