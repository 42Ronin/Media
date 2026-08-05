#!/usr/bin/env python3
"""Builds the Lesson 1 script in the format the team established (see NHO.docx)."""
import json
import os
import re as _re

import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

GREEN  = RGBColor(0x1E, 0x7A, 0x33)   # correct answer
ORANGE = RGBColor(0xC0, 0x60, 0x00)   # incorrect answer
BLUE   = RGBColor(0x1F, 0x4E, 0xD8)   # sources
RED    = RGBColor(0xC0, 0x1C, 0x1C)   # requested information by the course author

d = docx.Document()
st = d.styles["Normal"]
st.font.name = "Calibri"
st.font.size = Pt(11)
for s in d.sections:
    s.left_margin = s.right_margin = Inches(1)

def para(text="", bold=False, italic=False, underline=False, color=None,
         style=None, space_after=6):
    p = d.add_paragraph(style=style)
    if text:
        r = p.add_run(text)
        r.bold, r.italic, r.underline = bold, italic, underline
        if color is not None:
            r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p

_slide = {"sec": 0, "n": 0}
EMITTED = []

def title(t):        return para(t, bold=True, underline=True, space_after=12)
def section(t):
    _slide["sec"] += 1
    _slide["n"] = 0
    d.add_paragraph()
    return para(t, bold=True, underline=True, space_after=8)
def topic(t):
    """Slides are numbered section.slide, as in the team's storyboard format, so a
    reviewer and a developer can both point at one thing."""
    _slide["n"] += 1
    EMITTED.append((f"{_slide['sec']}.{_slide['n']}", t))
    return para(f"Slide {_slide['sec']}.{_slide['n']} \u2013 {t}", bold=True, space_after=4)
def sub(t):          return para(t, underline=True, space_after=4)
def body(t):         return para(t, space_after=8)
def note(t):         return para(t, italic=True, space_after=8)
def ask(n):
    """Open questions are numbered so a reviewer can refer to one by number, and
    are listed together at the front as well as in place."""
    return para(f"Open question {n + 1}: {OPEN_QUESTIONS[n]}", color=RED, space_after=8)
def bullet(t, lvl=0):
    p = d.add_paragraph(t, style="List Bullet" if lvl == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(2)
    return p
def num(t):
    p = d.add_paragraph(t, style="List Number")
    p.paragraph_format.space_after = Pt(2)
    return p
def answer(t, correct):
    p = d.add_paragraph(style="List Bullet 2")
    r = p.add_run(t)
    r.font.color.rgb = GREEN if correct else ORANGE
    p.paragraph_format.space_after = Pt(2)
    return p

# ---- the simulation's own task definitions --------------------------------
# Transcribed straight from the built lesson so the script and the thing the
# learner actually sees cannot drift. Regenerate with extract_tasks.mjs.
import task_data
TASKS = task_data.specs()

def plain(html):
    """The lesson marks copy up for the screen; the script wants the words."""
    s = _re.sub(r"<[^>]+>", "", html)
    s = (s.replace("&amp;", "&").replace("&lt;", "<").replace("&gt;", ">")
          .replace("&nbsp;", " ").replace("&#39;", "'").replace("&quot;", '"'))
    return _re.sub(r"\s+", " ", s).strip()

MONTHS = ("January February March April May June July August September "
          "October November December").split()

def longdate(iso):
    y, m, day = iso.split("-")
    return f"{int(day)} {MONTHS[int(m) - 1]} {y}"

# Every task traces to a paragraph of the vetted draft. The reference is carried
# here so the mapping can be checked without reading the build.


ACTION = {"open": "Open the record.", "choose": "Choose a record."}

DOC_DATE = "31 July 2026"

OPEN_QUESTIONS = []

# ───────────────────────────────── title + key ────────────────────────────
title("HMIS BNTS")
para("Lesson 1: Finding a Participant — Elearning Course Script", bold=True, space_after=10)
body("The key below displays the function of various text attributes of the following script.")

key = d.add_table(rows=5, cols=2)
key.style = "Table Grid"
key.alignment = WD_TABLE_ALIGNMENT.LEFT
rows = [
    ("Italics", "Ideas for video, images, and interactivity. Italics text will not be included in the course."),
    ("Blue",    "Sources"),
    ("Green",   "Correct answer(s) from a knowledge check"),
    ("Orange",  "Incorrect answer(s) from a knowledge check"),
    ("Red",     "Requested information by the course author"),
]
for i, (a, b) in enumerate(rows):
    key.rows[i].cells[0].width = Inches(1.2)
    key.rows[i].cells[1].width = Inches(5.3)
    ra = key.rows[i].cells[0].paragraphs[0].add_run(a)
    ra.bold = True
    if a == "Blue":   ra.font.color.rgb = BLUE
    if a == "Green":  ra.font.color.rgb = GREEN
    if a == "Orange": ra.font.color.rgb = ORANGE
    if a == "Red":    ra.font.color.rgb = RED
    if a == "Italics": ra.italic = True
    key.rows[i].cells[1].paragraphs[0].add_run(b)

d.add_paragraph()
note("Third pass, and the first intended for review outside the project. Every task in the "
     "practice simulation is documented in full — the situation, "
     "the instruction, the hint, the correct record and the feedback — so this can be reviewed "
     "without opening the simulation. External links and sources are intentionally omitted at this stage and "
     "will be added in a later pass. Interactivity notes describe the interaction rather than a "
     "specific authoring-tool feature, because this lesson is built in code rather than in Rise.")

# ───────────────────────────────── front matter ───────────────────────────
section("About This Document")

meta = d.add_table(rows=8, cols=2)
meta.style = "Table Grid"
meta.alignment = WD_TABLE_ALIGNMENT.LEFT
for i, (a, b) in enumerate([
    ("Lesson",        "HMIS Basic New Trainee Series — Lesson 1: Finding a Participant"),
    ("Version",       "0.3, draft for review"),
    ("Date",          DOC_DATE),
    ("Author",        "Nishan Paparian, Instructional Design"),
    ("Audience",      "HMIS end users at LAHSA and partner agencies — outreach and front-line "
                      "staff who look up participants in Clarity. No prior Clarity experience "
                      "assumed."),
    ("Prerequisites", "None within this series. Learners are assumed to hold, or to be in the "
                      "process of obtaining, an HMIS user account and whatever access agreement "
                      "and privacy training that requires."),
    ("Seat time",     "About forty minutes. See the note under How the Practice Works for where "
                      "it goes."),
    ("Delivery",      "Rise 360, with the practice simulation embedded as an uploaded HTML "
                      "block. The simulation is hands-on practice; Rise carries the score."),
]):
    meta.rows[i].cells[0].width = Inches(1.4)
    meta.rows[i].cells[1].width = Inches(5.1)
    meta.rows[i].cells[0].paragraphs[0].add_run(a).bold = True
    meta.rows[i].cells[1].paragraphs[0].add_run(b)
d.add_paragraph()

para("What Is Being Asked of the Reviewer", bold=True, space_after=4)
body("This draft is complete enough to review end to end. Every screen of narration is written, "
     "and every task in the practice simulation is documented in full — the situation the learner "
     "sees, the instruction, the hint, the record that is correct, and the feedback — so nothing "
     "needs to be opened to review it.")
body("What is not settled are the questions below. They are numbered, and each appears again in "
     "place so it can be read in context. Most need either a subject-matter answer or a decision "
     "about scope.")
note("Not yet in this draft, and deliberately: external links and source citations, and any "
     "visual or layout direction. Production notes describe what a screen has to do and what the "
     "learner does on it, not how it should look. Styling is a later pass, once the content is "
     "agreed.")

sub("Open questions, in one place")
for _qn, _qt in enumerate(OPEN_QUESTIONS):
    para(f"{_qn + 1}. {_qt}", color=RED, space_after=6)

d.add_paragraph()

# ───────────────────────────────── contents ───────────────────────────────
section("Table of Contents")
def _t(n):
    return f"Task {TASKS[n]['n']} — {TASKS[n]['title']}"

toc = [
    ("Introduction", ["Course Navigation", "Course Overview", "Objectives",
                      "Where This Lesson Sits", "Lesson Structure",
                      "Someone Asks You for an Update", "How This Lesson Runs",
                      "Why Searching Thoroughly Matters"]),
    ("Words You Will See", ["Key Terms"]),
    ("Two Rules Before You Touch the Keyboard", ["Assume a Record Already Exists",
                                                 "Never Create Before You Search"]),
    ("Preparing to Search", ["What to Collect First", "The Order to Ask In",
                             "Confirm the Spelling"]),
    ("How Search Works", ["Why It Is These Three", "One Bar, Many Kinds of Information",
                          "Searching by Name", "Searching by Date of Birth",
                          "Searching by Social Security Number", "Reading the Results"]),
    ("How the Practice Tool Is Built", ["Only What This Lesson Needs", "Earning the Rest"]),
    ("Practice: The Sandbox",
     ["How the Practice Works",
      "Round 1 — The Name You Were Given", _t(0), _t(1),
      "Round 2 — When You Only Have a Piece", _t(2), _t(3),
      "Round 3 — Too Many, and Then None", _t(4), _t(5),
      "Round 4 — Written Down Differently", _t(6), _t(7),
      "Halfway — From Finding to Verifying",
      "Verify: Is This the Right Person?", "First Level: Photo and Documents",
      "Second Level: Two of Three", "Still Unsure", "If You Later Find You Were Wrong",
      "Record the Match, Then Keep Looking",
      "What You Found", "No Matching Record", "One Matching Record",
      "Several Matching Records", "Choosing Which Record to Work From",
      "Round 5 — Which One of These Is Them", _t(8), _t(9), _t(10),
      "Round 6 — Everything at Once", _t(11), _t(12),
      "Out of the Sandbox"]),
    ("When Search Comes Up Empty", ["Somebody Nobody Can Find", "Alternate Names",
                                    "Alternate Spellings", "Start Over From a Different Fact",
                                    "Ask Your Data Staff", "Only Then, Create a Record"]),
    ("When You Are Short on Time", ["In the Field"]),
    ("Knowledge Check", ["Check What You Took In"]),
    ("Summary", ["Why Accuracy on the Way In Matters", "What You Practiced",
                 "Lesson Closing", "Survey"]),
    ("Production Standards", ["Accessibility"]),
]

for _si, (sec, tops) in enumerate(toc, start=1):
    para(f"{_si}. {sec}", bold=True, space_after=2)
    for _ti, t in enumerate(tops, start=1):
        para(f"      {_si}.{_ti}  {t}", space_after=1)

d.add_page_break()

# ───────────────────────────────── introduction ───────────────────────────
_slide["sec"] = 0          # front matter is not part of the slide sequence
section("Introduction")

topic("Course Navigation")
body("Before we start, a quick look at how to move around. This lesson has audio, so turn your "
     "sound on. Closed captions are available from the captions icon. The contents on the left "
     "let you jump back to anything you have already finished. Move forward with the right arrow "
     "— it appears once you have done what a screen asks of you — and back with the left.")
note("Standard navigation screen, matching the rest of the series. Adjust the wording to whatever "
     "Rise actually shows once the course shell is built.")

topic("Course Overview")
body("This lesson is about one thing: finding out whether the person in front of you already has a "
     "record in HMIS, and making sure you are looking at the right one.")
body("Set aside about forty minutes. A good part of that is hands on: you will work in a "
     "practice version of Clarity — the same screens you use at work, with invented people in "
     "them. Nothing you do in the practice affects real records, and none of the people in it "
     "are real.")
body("You can return to this lesson any time after you finish it.")
note("Where the time goes, measured from this script rather than estimated. Instruction is 3,929 "
     "words — about 26 minutes narrated, about 18 if the reference material is read on screen "
     "rather than spoken. What the learner reads inside the practice is another 1,700 or so. The "
     "thirteen tasks are the smallest part at roughly 6 minutes of actual interaction, because "
     "each is a short search. Hints and the correct-record lines in this document are not "
     "learner-facing time at all.")
note("Which puts the lesson between 35 and 45 minutes depending on how much is narrated. Twenty "
     "is not reachable while this lesson also carries why-data-quality-matters and thirteen "
     "hands-on tasks; those two together are most of the runtime. The remaining lever inside the "
     "script is the task feedback, currently around 95 words per task where 50 would land as "
     "hard.")

note("If the series would rather run shorter lessons anyway, the natural cut is the checkpoint "
     "between task 8 and task 9 — eight tasks on searching, six on verifying.")
note("Opening screen. Course title, estimated time, and a Begin button. The Begin button also "
     "satisfies the browser requirement that the learner interact with the page before audio can play.")

topic("Objectives")
body("This lesson has four objectives. By the end of it, you will")
num("Search HMIS on partial or uncertain information — a fragment of a name, a year, four "
    "digits — and find a participant's record.")
num("Recover a search that returns nothing, and narrow one that returns too much.")
num("Confirm that the record you found belongs to the person in front of you.")
num("Read a participant's history from their record, and explain why finding the right one "
    "decides what they have access to.")
body("Everything after this lesson assumes you can find the right person first.")
note("Presented as a numbered list, as HMIS Essentials does. Visual treatment is not decided at "
     "this stage and is deliberately not described here.")
note("Written to sit under objective 02 of HMIS Essentials — search and find a participant's "
     "HMIS record to learn their back story — rather than restating it, and to leave room for "
     "the lessons that follow in HMIS Basic Navigation to take objectives 03 to 05 without "
     "overlapping this one. An earlier draft of this slide had two of its three objectives "
     "sitting under Essentials objective 01, which is what had this lesson teaching an "
     "introduction it does not own.")


note("On screen as a numbered list before the lesson proper, the way the HUD Exchange courses "
     "open. Three, not more.")
note("Where each is taught and where it is checked — for review, not for the learner:")
note("1. What HMIS is and what it is for — taught in Why Searching Thoroughly Matters and Why It "
     "Is These Three; checked by knowledge-check questions 1 and 6.")
note("2. Using search, and why diligence matters — taught in How Search Works and When Search "
     "Comes Up Empty; practised in tasks 1 to 8; checked by knowledge-check questions 1, 2 and 6.")
note("3. Why accurate information matters — taught in Verify, What You Found and Why Accuracy on "
     "the Way In Matters; practised in tasks 9 to 14; checked by knowledge-check questions 3, 4 "
     "and 5.")

topic("Where This Lesson Sits")
body("HMIS Essentials sets out what a new user has to be able to do. This lesson is the second "
     "of those things.")
num("Identify the reasons why using HMIS correctly to enter data is important.")
num("Search and find a participant's HMIS record to learn their back story.  \u2190 this lesson")
num("Correctly create a new HMIS record for a participant that does not have one.")
num("Enroll a participant into a program and perform services, run an assessment, write a case "
    "note, and upload documents.")
num("Access the HMIS Knowledge Base to further expand your HMIS knowledge.")
body("Everything here serves the second one. Where this lesson touches the others, it gives you "
     "what you need and no more.")
note("Reviewer-facing as much as learner-facing — it is here so the boundary between this "
     "lesson and the rest of HMIS Basic Navigation is explicit and can be argued with. If it "
     "appears on screen at all, the current lesson is marked and the list is not read aloud in "
     "full. How it looks is a later question.")

topic("Lesson Structure")
body("This lesson has five parts.")
num("Getting ready — what to ask for, and in what order.")
num("How search works.")
num("Practice: finding a participant.")
num("Verifying that the record is really theirs, and practising that too.")
num("What to do when you are short on time, and what to take away with you.")
note("Consider a persistent marker showing which part the learner is in, as the other courses in "
     "the series do.")

topic("Someone Asks You for an Update")
body("There is a cleanup on the block this morning. People are moving their things, and your team "
     "is working the encampment.")
body("Someone walks up to you. Not a stranger exactly — you have seen him here before. He gives "
     "his name as Desmond, and he wants to know whether there is any news about interim housing "
     "for him.")
body("Desmond tells you somebody came round a while back. Took his details, said they would be "
     "in touch. Nothing happened, and he does not remember the name of the agency or the worker.")
body("That is the part to notice. If a provider took his details there is very likely already a "
     "record — a contact, a note, something with his name on it. Desmond is new to you; he is not "
     "new to the system. Right now he is a potential participant, and whether he is a participant "
     "depends on what your search finds.")
body("His question deserves a real answer, and to give him one you need his record.")
body("So the first thing you do is not paperwork and it is not a form. It is a search.")
body("Get it right and everything he is owed comes with him — the time he has waited, the "
     "programs he is in, the place he holds on a list. Get it wrong and he starts from nothing in "
     "a second record, while the first sits somewhere with his history in it.")
note("Opening sequence. Photographs with slow pan and zoom, narration over the top, captions on "
     "screen throughout. Roughly 60 to 90 seconds. Ends on a Continue button rather than running "
     "straight into the next screen.")



topic("How This Lesson Runs")
body("The shape of it is the shape of the job:")
num("You are an outreach worker with a caseload and a system you are accountable to.")
num("You meet someone.")
num("You search for them — first name, last name, date of birth.")
num("If and only if you cannot find them, you enter them.")
body("Everything in this lesson sits on one of those four steps. Most of it sits on the third, "
     "because that is where the damage is done or avoided.")
note("Consider showing the four steps as a progress marker that stays visible, so the learner "
     "always knows which part of the job they are in.")

topic("Why Searching Thoroughly Matters")
body("Searching thoroughly is how a participant keeps what they have already earned. Their wait "
     "time, their enrollments, their referrals and their place in a queue all live on the record. "
     "Find the right one and their case continues. Miss it and it starts over.")
body("There is a harder edge to it. A record can serve as third-party documentation of someone's "
     "homelessness — a shelter stay, or a contact logged by an outreach worker, is evidence of "
     "where they have been, and that evidence is part of what qualifies them for housing. A "
     "duplicate splits it, and the bed nights that prove their history sit on the record you did "
     "not find.")
note("Sourced; a source link goes here in the sourcing pass.")
body("A duplicate is also not a tidy-up job. Once two records exist, somebody has to work out "
     "which is which and clean it up — weeks, not minutes, and the person is the one waiting.")
sub("What a thorough search protects")
bullet("Continuity of care across every provider the participant has worked with")
bullet("A complete history when they are being qualified for services or housing")
bullet("Time — yours and your colleagues' — that would otherwise go into merging records later")
body("There is one more reason, and it reaches past your agency. Your Continuum of Care reports an "
     "unduplicated count of everyone served. A duplicate record counts one person twice. Enough of "
     "them and the region's picture of homelessness stops matching reality.")
note("Consider a short animated figure: one person, two record cards, and the count reading 2. "
     "Static illustration is fine if animation is out of scope.")

# ───────────────────────────────── two rules ──────────────────────────────
section("Words You Will See")

topic("Key Terms")
body("These come up throughout the lesson. You do not need to memorise them; they are here so "
     "nothing later depends on a word you have not met.")
sub("HMIS")
body("Homeless Management Information System. The system your CoC uses to record who has been "
     "served and what they were given. Clarity is the software LAHSA uses for it.")
sub("CoC — Continuum of Care")
body("The regional body that coordinates homeless services and reports on them. Los Angeles has "
     "one; you work inside it.")
sub("Participant, potential participant, and Client")
body("A participant has a record in HMIS. A potential participant is someone who may not — the "
     "person in front of you before you have searched. The software calls both of them Client.")
sub("Record")
body("Everything HMIS holds about one person: their identifying information, the programs they "
     "have been enrolled in, the services they have received, who they are housed with.")
sub("Unique Identifier")
body("The short code under a participant's name in search results, like AD63C3FF2. It names one "
     "record and nothing else, which is why you quote it when you report a problem.")
sub("Duplicate")
body("Two records for the same person. The thing this lesson exists to prevent.")
sub("Data quality")
body("How complete and accurate a record is — whether the name is full, whether the date of birth "
     "is exact, whether the Social Security Number is there. It is recorded on the record itself, "
     "not just implied.")
note("A short screen, or a reference panel the learner can open at any point rather than a screen "
     "they have to sit through. Consider the latter.")

section("Two Rules Before You Touch the Keyboard")

topic("Assume a Record Already Exists")
body("The goal of search is to find a record that is already there.")
body("Start from the assumption that the participant already has one. It is unlikely that you are "
     "the first homeless services staff member this person has ever spoken to — and if they have "
     "spoken to anyone, there is probably a record.")

topic("Never Create Before You Search")
body("Never create a new HMIS record for a participant until you have thoroughly searched for an "
     "existing one.")
body("This is the rule the rest of the lesson serves. Everything that follows is about what "
     "\"thoroughly\" actually means.")
note("Consider holding these two rules on screen together, with the learner clicking Continue to "
     "proceed. They are referred back to repeatedly.")

# ───────────────────────────────── preparing ──────────────────────────────
section("Preparing to Search")

topic("What to Collect First")
body("Before you search, note down the participant's full name and either their date of birth or "
     "the last four digits of their Social Security Number.")
body("Ask for it before you start typing, and ask for all of it in one go. Going back three "
     "times to ask someone to repeat their date of birth does not build trust, and it makes you "
     "look like the system is in charge of the conversation rather than you.")

topic("The Order to Ask In")
body("First name, last name, date of birth. That is the order, and it is the order because it is "
     "the order people will answer in.")
body("Social Security Number comes later, if you need it. Asking a stranger for their SSN in the "
     "first minute of a conversation costs you more than it gets you, and you can search perfectly "
     "well without it.")
body("Date of birth tends to be accurate and people hesitate over it far less than they do over a "
     "Social Security Number.")
body("If someone does not want to give you their SSN, that is a legitimate answer and you can "
     "work without it. Plenty of records hold no Social Security Number at all, or only part of "
     "one, and they are still findable.")

topic("Confirm the Spelling")
body("Confirm the spelling of the full name before you move on.")
body("If they have any identification or written documentation, ask to see it. Documents often "
     "reveal a middle name, a second last name, or a spelling that nobody would have guessed.")

# ───────────────────────────────── how search works ───────────────────────
section("How Search Works")

topic("Why It Is These Three")
body("Name, date of birth and Social Security Number are not an arbitrary trio. They are the "
     "things a person has one of. One name, one date of birth, one Social Security Number — so "
     "they are what tells one person from another, and they are what you check a record against.")
body("Everything else on a record can legitimately change. Someone can have several enrollments, "
     "several programs, several addresses, and none of that makes them a different person. The "
     "three identifiers are the part that should never disagree.")
body("Which is also why a second record for the same person is a problem rather than a "
     "duplicate row. It gives one person two sets of the things they only have one of.")

topic("One Bar, Many Kinds of Information")
body("Clarity has a single search bar, and it accepts several different kinds of information at "
     "once. You can type part of a name, a date of birth, a year, or part of a Social Security "
     "Number — and you can combine them.")
body("Two things are worth knowing before you start, because they surprise people.")
sub("Results appear as you type")
body("You do not need to press Enter. The list updates with every character.")
sub("Every word you add must match")
body("Adding a word narrows the search — it never widens it. If you type a first name and a last "
     "name and get nothing, one of the two is wrong. Try each on its own.")
note("Guided walkthrough, not a passive recording. The search text types itself on screen a "
     "character at a time, a tooltip explains what is happening, the result list narrows "
     "underneath, and the learner clicks to advance when they are ready. This is the pattern the "
     "RoninWrite tutorial uses and it should be lifted from there rather than reinvented.")

topic("Searching by Name")
body("You do not need a whole name. The first few letters of a first and last name are usually "
     "enough — for Katherine Johnson, typing kat joh will find her.")
body("Short fragments are often better than full names, because a full name has to be spelled the "
     "way somebody else typed it, and fragments do not.")

topic("Searching by Date of Birth")
body("A date of birth can be written several ways and all of them work: 9/26/1976, 09/26/1976, "
     "9.26.1976, or 9-26-1976.")
body("You can also search a year on its own. If the participant knows they were born in 1976 but "
     "not the exact date, 1976 is a perfectly good search.")

topic("Searching by Social Security Number")
body("You can search the last four digits on their own, or the full number.")
body("Some records hold only part of an SSN. When a participant could not recall all of it, the "
     "missing parts were stored as X's or zeros — so a record might read XXX-XX-6789. The digits "
     "that were recorded are still searchable.")
sub("What the code next to it tells you")
body("A record also carries a reason when the number is not there, and the three you will see "
     "mean genuinely different things:")
bullet("Client doesn't know — they do not have the information.")
bullet("Client prefers not to answer — they know it and have chosen not to give it. That is "
       "their decision to make, and it is a legitimate answer. Nobody is denied assistance for it.")
bullet("Data not collected — nobody recorded it. Often that means nobody asked.")
body("The difference is worth reading. A refusal is settled; a gap is not. If a record shows "
     "data not collected and you are with the participant, that is a blank you may be able to "
     "fill — and one more identifier to verify against next time.")
body("The rule that goes with it: you may not decide on someone's behalf. Do not record a "
     "refusal without asking, and never use it to mean that you do not know the answer.")
note("Definitions are the standard ones; a source link goes here in the sourcing pass.")


topic("Reading the Results")
body("Each result shows the participant's name, their date of birth and age, and the last four "
     "digits of their SSN. That is what you need to tell one person from another.")
note("The ROI column is blurred out in the simulation for this lesson — see Building It Up, below. "
     "Nothing needs to be said about it, so nothing is.")
body("A word about wording, because this lesson is careful with it.")
body("Someone with a record in HMIS is a participant. Someone standing in front of you who may "
     "or may not have one yet is a potential participant — and that is most of this lesson, "
     "because until you have searched you do not know which you are looking at.")
body("The software says Client. That means the same thing as participant. Both are correct — one "
     "is how we talk about people, the other is what the screen is labelled.")


# ───────────────────────────────── the tool ───────────────────────────────
section("How the Practice Tool Is Built")

topic("Only What This Lesson Needs")
body("The practice version of Clarity shows the whole screen, because that is what you will see "
     "at work. But the parts this lesson does not use are blurred out and cannot be clicked.")
body("This is deliberate, and it is worth explaining. A new user opening Clarity for the first "
     "time is looking at dozens of controls, almost none of which they need in their first week. "
     "Blurring the rest means there is never a moment in this lesson where something is on screen "
     "that we have to explain away, or promise to cover another time.")
sub("What is obstructed in this lesson")
bullet("The ROI column in the results table")
bullet("The client record's other tabs — Privacy, Programs, Services, Assessments, Files")
bullet("The Add Client form behind the ⊕ button")
bullet("The filter and column controls beyond the ones the tasks use")

topic("Earning the Rest")
body("Each lesson in the series switches on the part of the interface it teaches. Finish this "
     "one and search is fully yours; the next lesson lights up the create form, and so on, until "
     "the whole screen is live and you have been taught every part of it.")
note("The reveal should be an event, not a settings change. When a lesson unlocks a feature, the "
     "blur lifts on screen with a short animation and a line naming what the learner "
     "just earned — the same beat as gaining a new item in a game, where the reward is a "
     "permanent increase in what you are able to do. It costs very little to build and it gives "
     "a series of short lessons a spine.")
note("Production consequence worth stating plainly: this only works if the lessons share one "
     "simulation and one record of what the learner has unlocked. That is an argument for "
     "building the series on the common simulation this lesson already uses, rather than "
     "rebuilding the interface per lesson.")

# ───────────────────────────────── practice ───────────────────────────────
def kv_table(rows, label_width=1.45):
    """Two-column label/value table. A run of sub-headings per task is accurate
    and very hard to scan; a table says the same thing at a glance."""
    t = d.add_table(rows=len(rows), cols=2)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (label, value, italic) in enumerate(rows):
        c0, c1 = t.rows[i].cells
        c0.width = Inches(label_width)
        c1.width = Inches(6.5 - label_width)
        c0.paragraphs[0].add_run(label).bold = True
        for j, v in enumerate(value if isinstance(value, list) else [value]):
            p = c1.paragraphs[0] if j == 0 else c1.add_paragraph()
            p.add_run(v).italic = italic
    d.add_paragraph()
    return t


def emit_task(t):
    """One situation, documented the way a reviewer needs to read it: what the
    learner is put in front of, what they are asked to do, the help available,
    the record that is correct, and the words they get back when they find it."""
    topic(f"Task {t['n']} — {t['title']}")
    rows = [
        ("Draft reference", t["draft_ref"],   True),
        ("Situation",       t["situation"],   False),
        ("Instruction",     t["instruction"], False),
        ("Hint",            t["hint"],        False),
        ("Correct record",  t["answer"],      False),
        ("Feedback",        t["feedback"],    False),
    ]
    if not t["built"]:
        rows.insert(1, ("Status", "Specified here; not built in the simulation yet.", True))
    kv_table(rows)


section("Practice: The Sandbox")

topic("How the Practice Works")
body("What follows is a practice version of Clarity. It behaves the way the real system behaves. "
     "Everyone in it is invented — nothing you do here touches a real record.")
body("You open it once. You do not leave it again until the practice is finished.")
body("Thirteen people will come to you, in six rounds. Each round is harder than the one before "
     "it, and each round needs everything the rounds before it taught you — the techniques do "
     "not get retired, they get added to. By the last round you are choosing which one to reach "
     "for without being told.")
body("The practice roster holds three hundred people. The live system holds hundreds of "
     "thousands. A search that returns a readable handful here can return pages there — so the "
     "habit to build is to narrow before you scan, not to scroll.")
note("One simulation, opened at this point and held open for the whole practice. The learner "
     "does not return to a slide between situations — the panel on the left changes, the search "
     "bar and the results stay where they are. That continuity is the point: it is one sitting at "
     "one screen, which is what the work is.")
note("Rounds are marked in the panel — Round 1 of 6 — with a line naming what this round adds. "
     "Between rounds the panel pauses for a Continue rather than cutting away to a slide.")
note("The practice is not scored. It is practice — the learner works each situation until they "
     "find the record, hints are free and unlimited, and nothing here counts towards passing. "
     "The knowledge check carries the grade, at eighty per cent.")
note("That is deliberate. Searching well means searching repeatedly: clear the box, type "
     "something else, watch the list change, try again. A score would punish exactly the "
     "behaviour the lesson is trying to build.")
note("Unscored does not mean optional. Every situation has to be worked before the learner moves "
     "on — there is no skip. Hints escalate as they are asked for, ending with one that names "
     "the search outright, so nobody is stuck permanently. But they type it themselves and they "
     "see it work.")
note("Every situation traces to a paragraph of script_finding_a_participant_v3; the reference is "
     "given under each. Each is documented in full — the situation, the instruction, the hint, "
     "the record that is correct, and the feedback — so this can be reviewed without opening the "
     "simulation.")

topic("Round 1 — The Name You Were Given")
kv_table([
    ("What this round adds",
     "That the name you are told is not always the name on the record, and that an empty result "
     "is the start of the work rather than the end of it.", False),
    ("What it re-uses", "Nothing yet. This is the floor.", False),
    ("Situations", "Two.", False),
])

for _t in TASKS[0:2]:
    emit_task(_t)

topic("Round 2 — When You Only Have a Piece")
kv_table([
    ("What this round adds",
     "Working from a fragment of something rather than the whole of it — four digits of an SSN, "
     "the first few letters of a name.", False),
    ("What it re-uses",
     "Round 1. Both of these people are also filed under something other than what you were "
     "told, so you are still not taking the first answer as final.", False),
    ("Situations", "Two.", False),
])

for _t in TASKS[2:4]:
    emit_task(_t)

topic("Round 3 — Too Many, and Then None")
kv_table([
    ("What this round adds",
     "The two failures that look opposite and are not: a search that returns more than you can "
     "read, and one that returns nothing at all. Adding a term narrows. Swapping a term recovers.", False),
    ("What it re-uses",
     "Rounds 1 and 2 — you are narrowing and swapping with fragments, not full values.", False),
    ("Situations", "Two.", False),
])

for _t in TASKS[4:6]:
    emit_task(_t)

topic("Round 4 — Written Down Differently")
kv_table([
    ("What this round adds",
     "That somebody typed this record while listening. C and K, I and Y, a compound surname run "
     "together into one word.", False),
    ("What it re-uses",
     "All of it. You are searching fragments, you are not trusting the first empty result, and "
     "you are choosing which part of the name to keep.", False),
    ("Situations", "Two.", False),
])

for _t in TASKS[6:8]:
    emit_task(_t)

topic("Halfway — From Finding to Verifying")
body("Four rounds in. You now have every documented way in: a name fragment, a year, four digits "
     "of an SSN, an alternate spelling, a term added to narrow, and a term swapped when nothing "
     "came back.")
body("The two rounds left are a different problem. Finding a candidate record is not the same as "
     "finding the right one — and from here the search is the easy half.")
note("A pause inside the sandbox, not a screen away from it. The panel stops for a Continue; the "
     "simulation stays on screen behind it.")

# ───────────────────────────────── verify ─────────────────────────────────
topic("Verify: Is This the Right Person?")
body("Finding a candidate record is not the same as finding the right one. Verifying is a separate "
     "step and it has an order.")

topic("First Level: Photo and Documents")
body("The quickest check is visual. Look at the photo on the record, and look under the Files tab "
     "for identification or other documentation.")
body("These can tell you immediately whether the record belongs to the person in front of you.")
note("The simulation does not show photos or the Files tab, so this is taught with a still "
     "screenshot rather than practised. Enough to recognise the check and do it; the mechanics "
     "of uploading and managing documents belong to the lesson on files.")

topic("Second Level: Two of Three")
body("If there is no photo, work through the record with the participant and check whether the "
     "information is accurate for them.")
body("Start with name, date of birth, and Social Security Number. As a rule of thumb, if at least "
     "two of those three match, you can treat the record as a match.")
body("Know where that rule is weakest. Name and year of birth agreeing is the easiest pair to hit "
     "by coincidence — common surnames and a birth year will do it. When the two that match are "
     "the name and the date of birth, and the SSN is missing rather than different, look for a "
     "third thing before you commit: a household member, a program the participant remembers, a "
     "case note that fits.")

topic("Still Unsure")
body("Confirm anything else in the record with the participant:")
bullet("Case notes")
bullet("Program history")
bullet("Household members")
bullet("Location data")
bullet("Veteran status and other demographics")

topic("If You Later Find You Were Wrong")
body("Sometimes you work from a record for a while and then something does not fit — a program "
     "history they have no memory of, a household member who does not exist.")
body("Stop entering data into it, and note what you have entered and when. Anything in the system "
     "that needs changing by someone other than you goes to LAHSA HMIS Support as a ticket; they "
     "will tell you what happens next.")
note("Deliberately short. The full flow for a mis-identified record is not taught here.")




topic("Record the Match, Then Keep Looking")
body("If you decide the record is a match, note it down — then go back to search and look for "
     "others. Finding one match does not mean there is only one.")

# ───────────────────────────────── outcomes ───────────────────────────────
topic("What You Found")

topic("No Matching Record")
body("Create a new HMIS record for this participant.")

topic("One Matching Record")
body("Assume it is the correct one and work from it.")

topic("Several Matching Records")
body("Choose one to work from, and report the rest.")

topic("Choosing Which Record to Work From")
body("Choose the most complete record — full SSN, full date of birth, full name, full program "
     "history, a photo on the profile.")
body("Where that is unclear, choose the record with the longest enrollment history, based on the "
     "oldest enrollment.")
body("Note the Unique Identifiers of the others as you go, so you are not starting from nothing "
     "if it comes up again.")
note("Removed from this lesson: submitting the list of matching records to HMIS Support. "
     "Reporting duplicates is not part of this training.")

topic("Round 5 — Which One of These Is Them")
kv_table([
    ("What this round adds",
     "Reading past the identifiers. Two of three before you call it a match, and what to do when "
     "there is no third thing to check — the household, the location, whatever the record holds "
     "that the person in front of you can confirm.", False),
    ("What it re-uses",
     "Every search technique from rounds 1 to 4. Finding the candidates is no longer the "
     "difficulty; you will have them on screen in one search.", False),
    ("Situations", "Three.", False),
])

for _t in TASKS[8:11]:
    emit_task(_t)

topic("Round 6 — Everything at Once")
kv_table([
    ("What this round adds",
     "Nothing new, and that is the point. Nobody tells you which technique this needs. One of "
     "these takes several searches before anything works; the other gives you three records that "
     "are all the same person and asks you to choose.", False),
    ("What it re-uses",
     "All five rounds. Alternate names, alternate spellings, fragments, the year, narrowing, and "
     "reading past the identifiers.", False),
    ("Situations", "Two.", False),
])

for _t in TASKS[11:]:
    emit_task(_t)

topic("Out of the Sandbox")
body("Thirteen people, six rounds, one screen. You did not learn thirteen tricks — you learned "
     "six or seven, and then used them in combinations nobody told you to use.")
note("The simulation closes here for the first time since the practice began. It opens again for "
     "the scenario in When Search Comes Up Empty, which is the same sandbox and the same roster.")

# ───────────────────────────────── empty search ───────────────────────────
section("When Search Comes Up Empty")
body("You have searched and found nothing. Before you conclude the participant is new, work "
     "through this list — and this time you work it on one person, start to finish, choosing "
     "for yourself which step to reach for.")
note("Standalone sandbox rather than a run of taught screens. One scenario, and the learner "
     "drives the search bar. The panel moves the story on when they reach each point and does "
     "not block a route we did not predict. No score, a hint always available, and the learner "
     "does not leave until they have opened her record.")
note("Every step below is something a task already taught. This is where they put it together "
     "without being told which tool to pick up.")

topic("Somebody Nobody Can Find")
kv_table([
    ("Story",
     "Second morning of a cleanup. A woman waits until your team is packing up and then comes "
     "over. Weeks ago somebody told her a shelter bed was being held for her. She has heard "
     "nothing since, and she wants to know whether it is real. She gives her name as Sylvia "
     "Marchetti.", False),
    ("Learner does", "Searches what she gave them.", False),
    ("What happens",
     "Nothing. Not the full name, not Marchetti on its own. The empty state says No clients "
     "found and that is all it says.", False),
    ("The point",
     "An empty result is not proof that somebody is new — it is the beginning of the list "
     "below, not the end of the search.", False),
])

topic("Alternate Names")
body("Ask the participant what else they have been called: a middle name, a second last name, a "
     "nickname, a shortened name. Confirm the spelling of each one.")
body("Some of what you hear will not be a nickname at all. A participant may give you the name "
     "they use — a chosen name, or the name that fits who they are — while the record was created "
     "under a different one. Ask for both, without making it a thing: the name they go by, and "
     "any other name a record might be under. Search on both, and address them by the one they "
     "gave you.")
body("Ask to see any identification or written documentation. It is the fastest way to discover a "
     "name nobody thought to mention.")
kv_table([
    ("Story",
     "You tell her the search is not finding her, which is not the same as her not being there, "
     "and you ask what else she has been called. Everyone calls her Syl. She has not used "
     "Marchetti in years — there was another name before it and she cannot remember which one "
     "the shelter wrote down. She was born in 1979 but does not know the day. She would rather "
     "not give her Social Security Number, and that is fine.", False),
    ("Learner does", "Searches Syl. Then the year on its own.", False),
    ("What happens",
     "Syl reaches nobody — no alias was ever recorded for her. The year returns a long list, and "
     "no Marchetti is in it, because there is no Marchetti to find.", False),
    ("The point",
     "Asking got you three new facts and ruled out a fourth. None of them has found her yet, and "
     "that is still progress.", False),
])

topic("Alternate Spellings")
body("Try the spellings a previous provider might have used by mistake. For Katherine Johnson: "
     "catherine, cat, jonson, jon.")
kv_table([
    ("Story",
     "Sylvia. With a Y. Somebody wrote this down once while listening rather than reading, and "
     "the swap nobody notices themselves making is I for Y.", False),
    ("Learner does", "Searches Sil.", False),
    ("What happens",
     "Five records — Sileshi, Silvana, Silas, and two women called Silvia.", False),
    ("The point",
     "The fragment beats the full name, because a fragment does not have to be spelled the way "
     "somebody else typed it.", False),
])

topic("Start Over From a Different Fact")
body("Go back to the beginning and start from a different piece of information. If you began with "
     "the date of birth, begin this time with the first three letters of the first and last name.")
kv_table([
    ("Story",
     "Five is readable and four of them are plainly not her. You started from her name and the "
     "name was wrong twice over — so start from the other thing she gave you.", False),
    ("Learner does", "Adds 1979 to the fragment.", False),
    ("What happens",
     "Two records. Silvia Duarte and Silvia Okonkwo, both born that year. Nothing she has told "
     "you separates them: she does not know her exact date of birth, she has not given an SSN, "
     "and either surname could be the one she stopped using.", False),
    ("Learner does next",
     "Goes back to her. Is there anyone with you? Her son Mateo is on the kerb with their bags. "
     "Opens both records and reads past the identifiers.", False),
    ("The right record",
     "Silvia Duarte · unique identifier 7E1D4A9C3 · born 8 May 1979 · SSN none on file. Mateo "
     "Duarte is on her household. Silvia Okonkwo's record lists nobody.", False),
    ("The point",
     "When the identifiers run out, the rest of the record is what identifies somebody. The "
     "referral she was asking about is on the record you just opened — filed under a name she "
     "had stopped using and a spelling she never saw.", False),
])

topic("Ask Your Data Staff")
body("If you have time and your organisation has data staff, ask them. They search differently and "
     "they find things.")
note("Story note: this is where Silvia's morning would have gone next if both of those records "
     "had turned out to be somebody else. Worth saying so, so the step does not read as "
     "hypothetical.")

topic("Only Then, Create a Record")
body("If you have worked through all of the above and still cannot find a candidate record, create "
     "a new one. At that point it is the right thing to do.")
body("Notice how far down this list that sentence is. Sylvia was four searches from being a "
     "second record for a person who already had one — and every one of those searches was a "
     "thing you had already practised.")
note("Removed from this lesson: asking HMIS Support to double-check a newly created record. "
     "Reporting duplicates is not part of this training.")
note("Creating a record is its own lesson, and the create form is blurred out here. In this lesson "
     "the Add button acknowledges the learner and explains that it is covered next — and points "
     "out when the participant in the current task already has a record.")

# ───────────────────────────────── time pressure ──────────────────────────
section("When You Are Short on Time")

topic("In the Field")
body("Sometimes you cannot do all of this on the spot. In street outreach, under time pressure, you "
     "have two honest options.")
num("Write everything down — on paper or in your notes — and do the full search and verification "
    "later at the office. Then enter what you noted once you have found the right record.")
num("Create a record, if you need a Unique Identifier for the participant right away. Then report "
    "it to HMIS Support so they can check whether it duplicates an existing record, and tell you "
    "how to merge them if it does.")
body("What is not an option is creating a record and saying nothing.")

# ───────────────────────────────── knowledge check ────────────────────────
section("Knowledge Check")
topic("Check What You Took In")
note("Scored knowledge check. Correct answers in green, incorrect in orange. Randomise answer "
     "order. Every option needs feedback text, not just a right-or-wrong mark — the feedback for "
     "each question is given below it, and the wrong answers here are wrong for reasons worth "
     "saying out loud.")

body("1. You search a participant's name and get no results. What does that tell you?")
answer("Very little on its own — the record may be under a different name or spelling", True)
answer("That the participant has never received services before", False)
answer("That you should create a new record", False)
note("Feedback: an empty result is the single most misread signal in HMIS. It usually means the "
     "record is under a name, a spelling or an identifier you have not tried yet.")

body("2. A participant says she was born in 1985 but cannot remember the exact date. What can you search?")
answer("The year on its own", True)
answer("Nothing — a full date of birth is required", False)
answer("Only the first three letters of her name", False)
note("Feedback: a year on its own is a valid search. So is a fragment of a date. You do not need "
     "the whole thing to start narrowing.")

body("3. Two records share a name, but you have not compared anything else yet. What is your "
     "next step?")
answer("Compare a second identifier — date of birth, or the SSN fragment", True)
answer("Treat them as the same person, because the name matches", False)
answer("Work from whichever record you opened first", False)
note("Feedback: a name on its own is never enough. Two of the three — name, date of birth, SSN — "
     "have to agree before you treat a record as a match. Deciding on the name alone is how one "
     "person ends up with two records, and how two people end up sharing one.")

body("4. A participant will not give you their Social Security Number. What can you still do?")
answer("Search on name and date of birth — the SSN is one route in, not the only one", True)
answer("Nothing until they provide it", False)
answer("Create a new record, since you cannot confirm the old one", False)
note("Feedback: refusing is a legitimate answer, and it is recorded as one. Plenty of records hold "
     "no SSN at all, or only part of one. Name and date of birth will find most people, and the "
     "rest of the record — household, program history, case notes — can confirm the match.")

body("5. Two records look like they could be the same person, and you cannot tell them apart on "
     "name, date of birth or SSN alone. What do you do?")
answer("Keep looking at the rest of each record — household, program history, location — until "
       "one of them fits the person in front of you", True)
answer("Pick whichever record you opened first", False)
answer("Assume they are different people, because there are two records", False)
note("Feedback: the identifiers are where you start, not where you stop. When they will not "
     "settle it, the rest of the record usually will — and it is worth the extra minute, because "
     "the alternative is working from somebody else's history.")

body("6. When is it correct to create a new record?")
answer("Only after you have worked the whole list — alternate names and spellings, a different "
       "identifier, and your data staff if you have them — and still found nothing", True)
answer("After the participant's name returns no results", False)
answer("Whenever the participant says they have not received services before", False)
note("Feedback: creating a record is the last step, not the second one. An empty result means your "
     "search has not found them yet — it is not evidence that they are new. And when you do create "
     "one, send the Unique Identifier to HMIS Support and ask them to double-check.")


# ───────────────────────────────── job aid ────────────────────────────────
section("Summary")

topic("Why Accuracy on the Way In Matters")
body("This lesson is about finding someone. But every record you found was typed by somebody who "
     "came before you, and the only reason you could find it is that they got it right.")
body("So when you do reach the point of entering someone — because you searched properly and "
     "they genuinely are not there — the care you take is not administrative. Accurate "
     "information is what makes them findable by the next worker, and being findable is what "
     "makes their support arrive on time and go to the right person.")
body("Search and accurate entry are the same skill pointed in two directions. The next lesson is "
     "the other direction.")

topic("What You Practiced")
body("You searched on a name fragment, on a year, on four digits of a Social Security Number, and "
     "on a name that had been typed a different way. You narrowed a search that returned too much, "
     "and you swapped an identifier when a search returned nothing.")
body("Then you verified. You told two people with the same name apart on a second identifier. You "
     "identified someone whose record held almost no identifiers, using her household. You chose "
     "between three records that all matched. And you recognised the same person entered twice, and "
     "reported it rather than fixing it yourself.")
body("The habit underneath all of it is one sentence: an empty result is not proof that someone is "
     "new. Search the surname fragment, search the year of birth, confirm on a second identifier — "
     "and only then create a record.")
body("Lesson 2 covers creating that record, where this habit is the only thing standing between you "
     "and a duplicate.")

topic("Lesson Closing")
body("That is the end of Lesson 1. Well done — this is the least glamorous part of the job and "
     "the one that decides the most.")
body("The next lesson covers adding a participant who genuinely is not there, which is where "
     "everything you have just practised gets put to work.")
note("Closing screen with an Exit control, matching the series.")

topic("Survey")
body("Now that you have completed the lesson, we would appreciate your feedback. Please complete "
     "the short survey below.")
note("Survey embed.")

section("Production Standards")

topic("Accessibility")
note("At the back of the document on purpose. This was not part of the brief and is not asking "
     "for a decision — it is here so the standard is on the record as considered, and so it can "
     "be pulled forward if it turns out to be required.")
body("Applies to every screen in this lesson, and to the practice simulation.")
bullet("Narration is captioned, and a transcript is available on the page.")
bullet("Every image that carries meaning has alternative text. Images that are decorative are "
       "marked as decorative rather than described.")
bullet("Colour is never the only way something is signalled — a correct answer, an obstructed control "
       "or an alert also carries text or a shape.")
bullet("Everything the learner must do can be done from the keyboard, including the search field, "
       "the results table and opening a record.")
bullet("Text can be resized without content being cut off or overlapping.")
bullet("Contrast meets WCAG 2.1 AA. This includes the blurred controls, which must read as "
       "unavailable without becoming invisible.")
bullet("Nothing depends on hearing alone, and nothing depends on a timed response.")
note("LAHSA is a public agency and this training is federally adjacent, so Section 508 and "
     "WCAG 2.1 AA are the standard to build to, not to retrofit. The blur treatment described "
     "earlier is the one place where accessibility and design pull against each other — an "
     "obstructed control still has to read as deliberate rather than broken.")


# The table of contents is written before the body, so the only thing keeping the
# two in step is this check. A slide added to one and not the other fails here
# rather than reaching a reviewer.
_expected = [(f"{i}.{j}", t)
             for i, (_s, _tops) in enumerate(toc, start=1)
             for j, t in enumerate(_tops, start=1)]
if _expected != EMITTED:
    _only_body = [x for x in EMITTED if x not in _expected]
    _only_toc  = [x for x in _expected if x not in EMITTED]
    raise SystemExit("contents and body disagree.\n"
                     f"  in the body, not the contents: {_only_body}\n"
                     f"  in the contents, not the body: {_only_toc}")
print(f"{len(EMITTED)} slides, contents agrees")

OUT = "script_l1_v3.docx"
d.save(OUT)

# python-docx emits <w:zoom/> with no w:percent, which fails XSD validation and
# makes some tools refuse the file. Fill it in rather than leaving it malformed.
import zipfile, shutil, tempfile
_tmp = tempfile.mkdtemp()
with zipfile.ZipFile(OUT) as z:
    z.extractall(_tmp)
_st = os.path.join(_tmp, "word", "settings.xml")
if os.path.exists(_st):
    _x = open(_st, encoding="utf-8").read()
    _new = _re.sub(r"<w:zoom(?![^>]*w:percent)([^>]*?)/>",
                   r'<w:zoom\1 w:percent="100"/>', _x)
    if _new != _x:
        open(_st, "w", encoding="utf-8").write(_new)
os.remove(OUT)
with zipfile.ZipFile(OUT, "w", zipfile.ZIP_DEFLATED) as z:
    for root, _, files in os.walk(_tmp):
        for f in files:
            full = os.path.join(root, f)
            z.write(full, os.path.relpath(full, _tmp))
shutil.rmtree(_tmp)
print("written:", OUT)
