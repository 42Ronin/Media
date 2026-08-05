#!/usr/bin/env python3
"""Cut one section out of the built script, for pasting into a draft someone is
editing by hand.

    python3 snippet_section.py "When Search Comes Up Empty"

Copies the real elements — paragraphs, tables, formatting — so what lands in the
draft is exactly what the script says, not a re-rendering of it.
"""
import copy
import os
import re
import shutil
import sys
import tempfile
import zipfile

from docx import Document

HERE = os.path.dirname(os.path.abspath(__file__))
WANTED = sys.argv[1] if len(sys.argv) > 1 else "When Search Comes Up Empty"

src = Document(os.path.join(HERE, "script_l1_v3.docx"))


def is_heading(p):
    """Section headings are the only paragraphs that are bold and underlined."""
    rs = [r for r in p.runs if r.text.strip()]
    return bool(rs) and rs[0].bold and rs[0].underline


heads = [p for p in src.paragraphs if is_heading(p)]
names = [p.text.strip() for p in heads]
if WANTED not in names:
    raise SystemExit(f"no section called {WANTED!r}. Try one of:\n  " + "\n  ".join(names))
k = names.index(WANTED)

items = list(src.element.body)
i = items.index(heads[k]._p)
j = items.index(heads[k + 1]._p) if k + 1 < len(heads) else len(items) - 1

out = Document()
out.styles["Normal"].font.name = src.styles["Normal"].font.name
out.styles["Normal"].font.size = src.styles["Normal"].font.size
for el in items[i:j]:
    out.element.body.insert(len(out.element.body) - 1, copy.deepcopy(el))

slug = re.sub(r"[^a-z0-9]+", "-", WANTED.lower()).strip("-")
OUT = os.path.join(HERE, f"section-{slug}.docx")
out.save(OUT)

# python-docx writes <w:zoom> without the w:percent the schema requires
tmp = tempfile.mkdtemp()
with zipfile.ZipFile(OUT) as z:
    z.extractall(tmp)
stg = os.path.join(tmp, "word", "settings.xml")
if os.path.exists(stg):
    x = open(stg, encoding="utf-8").read()
    nx = re.sub(r"<w:zoom(?![^>]*w:percent)([^>]*?)/>", r'<w:zoom\1 w:percent="100"/>', x)
    if nx != x:
        open(stg, "w", encoding="utf-8").write(nx)
os.remove(OUT)
with zipfile.ZipFile(OUT, "w", zipfile.ZIP_DEFLATED) as z:
    for root, _, files in os.walk(tmp):
        for f in files:
            full = os.path.join(root, f)
            z.write(full, os.path.relpath(full, tmp))
shutil.rmtree(tmp)

d = Document(OUT)
slides = [p.text.strip() for p in d.paragraphs if p.text.strip().startswith("Slide ")]
print(f"wrote {os.path.basename(OUT)} — {len(slides)} slides, {len(d.tables)} tables")
