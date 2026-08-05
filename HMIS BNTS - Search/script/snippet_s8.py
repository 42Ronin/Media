#!/usr/bin/env python3
"""Section 8 only — tasks 1-8 — as one wide table: a row per task, a column per
field, so the set can be scanned by scrolling rather than read in sequence.
Landscape, because six columns of prose need the width. Paste into a draft that
is being edited by hand; this is not a replacement for the script."""
import json, os, re
import docx
from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
TASKS = json.load(open(os.path.join(HERE, "tasks.json")))["tasks"]

d = docx.Document()
st = d.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(11)
for s in d.sections:                       # landscape: six columns of prose
    s.orientation = WD_ORIENT.LANDSCAPE
    s.page_width, s.page_height = s.page_height, s.page_width
    s.left_margin = s.right_margin = Inches(0.6)
    s.top_margin = s.bottom_margin = Inches(0.6)

def para(text="", bold=False, italic=False, space_after=6):
    p = d.add_paragraph(); r = p.add_run(text)
    r.bold, r.italic = bold, italic
    p.paragraph_format.space_after = Pt(space_after)
    return p

def plain(html):
    s = re.sub(r"<[^>]+>", "", html)
    s = (s.replace("&amp;", "&").replace("&lt;", "<").replace("&gt;", ">")
          .replace("&nbsp;", " ").replace("&#39;", "'").replace("&quot;", '"'))
    return re.sub(r"\s+", " ", s).strip()

MONTHS = ("January February March April May June July August September "
          "October November December").split()
def longdate(iso):
    y, m, day = iso.split("-"); return f"{int(day)} {MONTHS[int(m)-1]} {y}"

DRAFT_REF = {
    "nickname":  "alternate names and nicknames",
    "year":      "search by year of birth alone",
    "last4":     "search by the last four digits of the Social Security Number",
    "fragments": "search the first letters of the first and last name, as in kat joh",
    "narrow":    "narrowing an over-broad result set by adding a term",
    "swap":      "replacing one search term with another piece of information",
    "spelling":  "alternate spellings, C and K, I and Y",
    "surname":   "second last names and compound surnames",
}
HEADINGS = {
    "nickname":  "The Name He Gave You",
    "year":      "When the Last Name May Have Changed",
    "last4":     "Starting From the Last Four of the SSN",
    "fragments": "A Name You Cannot Spell",
    "narrow":    "Too Many Results",
    "swap":      "Nothing Comes Back",
    "spelling":  "Spelled the Way Someone Else Heard It",
    "surname":   "A Surname Typed as One Word",
}


COLS = [
    ("Task",           1.35),
    ("Situation",      2.10),
    ("Instruction",    1.05),
    ("Hint",           1.75),
    ("Correct record", 1.30),
    ("Feedback",       2.15),
]

def cell(c, text, bold=False, italic=False, size=9):
    p = c.paragraphs[0] if not c.paragraphs[0].text else c.add_paragraph()
    r = p.add_run(text); r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(2)

t = d.add_table(rows=1, cols=len(COLS))
t.style = "Table Grid"
t.alignment = WD_TABLE_ALIGNMENT.LEFT
t.autofit = False

# repeat the header when the table runs over a page
hdr = t.rows[0]
tr = hdr._tr.get_or_add_trPr()
el = OxmlElement("w:tblHeader"); el.set(qn("w:val"), "true"); tr.append(el)
for i, (name, w) in enumerate(COLS):
    hdr.cells[i].width = Inches(w)
    cell(hdr.cells[i], name, bold=True)

for n, task in enumerate(TASKS[:8], start=2):
    row = t.add_row()
    for i, (_name, w) in enumerate(COLS):
        row.cells[i].width = Inches(w)
    bits = [task["answerName"], task["answerId"]]
    if task["answerDob"]:
        bits.append(longdate(task["answerDob"]))
    bits.append("SSN " + (task["answerSsn"] or "none on file"))

    cell(row.cells[0], f"Slide 8.{n}", size=8)
    cell(row.cells[0], f"Task {task['n']} — {HEADINGS[task['id']]}", bold=True)
    cell(row.cells[0], DRAFT_REF[task["id"]] + ".", italic=True, size=8)
    cell(row.cells[1], plain(task["brief"]))
    cell(row.cells[2], plain(task["ask"]))
    cell(row.cells[3], plain(task["hint"]))
    for b in bits:
        cell(row.cells[4], b)
    cell(row.cells[5], plain(task["teach"]))

d.add_paragraph()
para("Slide 8.10 – Checkpoint: From Finding to Verifying", bold=True, space_after=4)
para("That is the search half. You now have every documented way in: a name fragment, a year, "
     "four digits of an SSN, an alternate spelling, a term added to narrow, and a term swapped "
     "when nothing came back.")
para("Finding a candidate record is not the same as finding the right one. The rest of this "
     "lesson is verifying: proving a record belongs to the person in front of you, and deciding "
     "what to do when more than one does.")

OUT = os.path.join(HERE, "section-8-tasks-as-columns.docx")
d.save(OUT)
import zipfile, shutil, tempfile
tmp = tempfile.mkdtemp()
with zipfile.ZipFile(OUT) as z: z.extractall(tmp)
stg = os.path.join(tmp, "word", "settings.xml")
if os.path.exists(stg):
    x = open(stg, encoding="utf-8").read()
    nx = re.sub(r"<w:zoom(?![^>]*w:percent)([^>]*?)/>", r'<w:zoom\1 w:percent="100"/>', x)
    if nx != x: open(stg, "w", encoding="utf-8").write(nx)
os.remove(OUT)
with zipfile.ZipFile(OUT, "w", zipfile.ZIP_DEFLATED) as z:
    for root, _, files in os.walk(tmp):
        for f in files:
            full = os.path.join(root, f); z.write(full, os.path.relpath(full, tmp))
shutil.rmtree(tmp)
print("wrote", os.path.basename(OUT))
