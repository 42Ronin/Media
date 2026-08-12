#!/usr/bin/env python3
"""Cut one section's tasks as a single wide table — a row per task, a column per
field — for pasting into a draft that is being edited by hand.

    python3 snippet_tasks.py 8      # the eight search tasks
    python3 snippet_tasks.py 11     # the five verification tasks

Landscape, because six columns of prose need the width. Copy reads from
task_data, the same source the script uses, so a snippet cannot drift from it.
This is not a replacement for the script.
"""
import os
import re
import shutil
import sys
import tempfile
import zipfile

import docx
from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import task_data

HERE = os.path.dirname(os.path.abspath(__file__))

SECTIONS = {
    "8":  ("Section 8 — Practice: Finding a Participant", task_data.SEARCH_IDS, 2),
    "11": ("Section 11 — Practice: Verifying a Record",   task_data.VERIFY_IDS, 1),
}

which = sys.argv[1] if len(sys.argv) > 1 else "8"
if which not in SECTIONS:
    raise SystemExit(f"pick a section: {', '.join(SECTIONS)}")
heading, ids, first_slide = SECTIONS[which]

d = docx.Document()
st = d.styles["Normal"]
st.font.name = "Calibri"
st.font.size = Pt(11)
for s in d.sections:
    s.orientation = WD_ORIENT.LANDSCAPE
    s.page_width, s.page_height = s.page_height, s.page_width
    s.left_margin = s.right_margin = Inches(0.6)
    s.top_margin = s.bottom_margin = Inches(0.6)

COLS = [("Task", 1.35), ("Situation", 2.10), ("Instruction", 1.05),
        ("Hint", 1.75), ("Correct record", 1.30), ("Feedback", 2.15)]


def cell(c, text, bold=False, italic=False, size=9):
    p = c.paragraphs[0] if not c.paragraphs[0].text else c.add_paragraph()
    r = p.add_run(text)
    r.bold, r.italic = bold, italic
    r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(2)


def fill(c, value):
    for v in (value if isinstance(value, list) else [value]):
        cell(c, v)


t = d.add_table(rows=1, cols=len(COLS))
t.style = "Table Grid"
t.alignment = WD_TABLE_ALIGNMENT.LEFT
t.autofit = False

hdr = t.rows[0]
el = OxmlElement("w:tblHeader")            # repeat the header across pages
el.set(qn("w:val"), "true")
hdr._tr.get_or_add_trPr().append(el)
for i, (name, w) in enumerate(COLS):
    hdr.cells[i].width = Inches(w)
    cell(hdr.cells[i], name, bold=True)

specs = {s["id"]: s for s in task_data.specs()}
for n, tid in enumerate(ids, start=first_slide):
    spec = specs[tid]
    row = t.add_row()
    for i, (_n, w) in enumerate(COLS):
        row.cells[i].width = Inches(w)
    cell(row.cells[0], f"Slide {which}.{n}", size=8)
    cell(row.cells[0], f"Task {spec['n']} — {spec['title']}", bold=True)
    cell(row.cells[0], spec["draft_ref"], italic=True, size=8)
    if not spec["built"]:
        cell(row.cells[0], "Not built yet.", italic=True, size=8)
    fill(row.cells[1], spec["situation"])
    fill(row.cells[2], spec["instruction"])
    fill(row.cells[3], spec["hint"])
    fill(row.cells[4], spec["answer"])
    fill(row.cells[5], spec["feedback"])

OUT = os.path.join(HERE, f"section-{which}-tasks-as-columns.docx")
d.save(OUT)

# python-docx writes <w:zoom> without the w:percent the schema requires
tmp = tempfile.mkdtemp()
with zipfile.ZipFile(OUT) as z:
    z.extractall(tmp)
stg = os.path.join(tmp, "word", "settings.xml")
if os.path.exists(stg):
    x = open(stg, encoding="utf-8").read()
    nx = re.sub(r"<w:zoom(?![^>]*w:percent)([^>]*?)/>", r'<w:zoom\1 w:percent="100"/>', x)
    if nx != x:
        open(stg, "w", encoding="utf-8").write(nx)
os.remove(OUT)
with zipfile.ZipFile(OUT, "w", zipfile.ZIP_DEFLATED) as z:
    for root, _, files in os.walk(tmp):
        for f in files:
            full = os.path.join(root, f)
            z.write(full, os.path.relpath(full, tmp))
shutil.rmtree(tmp)

print(f"wrote {os.path.basename(OUT)} — {len(ids)} tasks")
for tid, why in task_data.debts().items():
    if tid in ids:
        print(f"  note: '{tid}' copy differs from the build. {why}")
