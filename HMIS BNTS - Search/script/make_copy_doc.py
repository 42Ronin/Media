#!/usr/bin/env python3
"""Render every learner-facing line from the HTML tools as an editable script.

    node script/extract_copy.mjs && python3 script/make_copy_doc.py

Out comes `Learner Text - HMIS BNTS Search (Sections 7, 10, 11).docx`: the coaching
copy from sections 7, 10 and 11 and the four step embeds, one editable block at a
time, each with a reference so an edit can be carried back to the exact string.

This is a WORKING document, not a generated report — the point of it is to be edited.
The build does not read it back; whoever edits it hands the edits over and they are
applied to `src/lesson1.template.html`, which stays the source of truth. Re-running
this overwrites the file, so do not edit in place and re-run.

The simulated product's own wording is deliberately absent. Menu items, column
headings, field labels and the empty state are Clarity's, not ours, and are not
ours to rewrite.
"""
import html as _html
import json
import os
import re

import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = os.path.dirname(os.path.abspath(__file__))
DATA = json.load(open(os.path.join(HERE, "copy.json"), encoding="utf-8"))
OUT = os.path.join(HERE, "..", "Learner Text - HMIS BNTS Search.docx")

TEAL   = RGBColor(0x06, 0x68, 0x88)
GREY   = RGBColor(0x5F, 0x77, 0x87)
GREEN  = RGBColor(0x1E, 0x7A, 0x33)
ORANGE = RGBColor(0xC0, 0x60, 0x00)
RED    = RGBColor(0xC0, 0x1C, 0x1C)

d = docx.Document()
st = d.styles["Normal"]
st.font.name = "Calibri"
st.font.size = Pt(11)
for s in d.sections:
    s.left_margin = s.right_margin = Inches(0.9)


def para(text="", bold=False, italic=False, underline=False, color=None,
         size=None, space_after=6, align=None):
    p = d.add_paragraph()
    if text:
        r = p.add_run(text)
        r.bold, r.italic, r.underline = bold, italic, underline
        if color is not None:
            r.font.color.rgb = color
        if size is not None:
            r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    return p


def rich(markup, space_after=10, indent=0.25):
    """One block of copy, with the emphasis the learner sees kept.

    `<b>` is what the copy stresses and `<code>` is what they are told to type —
    both change how a line reads aloud, so both survive into the document. Every
    other tag is layout and goes. Paragraph breaks are kept: several of these
    strings are two or three paragraphs and read as one wall without them."""
    for chunk in re.split(r"</p>\s*<p[^>]*>|<br\s*/?>", markup or ""):
        chunk = re.sub(r"</?p[^>]*>", "", chunk).strip()
        if not chunk:
            continue
        p = d.add_paragraph()
        p.paragraph_format.left_indent = Inches(indent)
        p.paragraph_format.space_after = Pt(space_after)
        for piece in re.split(r"(<b>.*?</b>|<code>.*?</code>|<span class=\"q\">.*?</span>)",
                              chunk, flags=re.S):
            if not piece:
                continue
            m = re.match(r"<b>(.*?)</b>$", piece, re.S)
            if m:
                r = p.add_run(_plain(m.group(1))); r.bold = True; continue
            m = re.match(r"<code>(.*?)</code>$", piece, re.S)
            if m:
                r = p.add_run(_plain(m.group(1)))
                r.font.name = "Consolas"; r.font.color.rgb = TEAL; continue
            m = re.match(r"<span class=\"q\">(.*?)</span>$", piece, re.S)
            if m:
                r = p.add_run(_plain(m.group(1))); r.italic = True; continue
            t = _plain(piece)
            if t:
                p.add_run(t)
    return


def _plain(s):
    return _html.unescape(re.sub(r"<[^>]+>", "", s or "")).replace("\n", " ")


def ref(text):
    """The handle for an edit. Every block carries one so a rewrite can be pointed
    at the exact string it replaces."""
    return para(text, size=8, color=GREY, space_after=2)


def label(text, color=None):
    return para(text, bold=True, size=9, color=color or TEAL, space_after=2)


def h1(text):
    d.add_page_break()
    return para(text, bold=True, size=17, color=TEAL, space_after=10)


def h2(text, space_after=8):
    return para(text, bold=True, size=13, space_after=space_after)


def note(text):
    return para(text, italic=True, color=GREY, space_after=10)


def rule():
    p = d.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.add_run("—" * 46).font.color.rgb = RGBColor(0xD0, 0xDA, 0xE0)
    return p


def block(reference, kind, markup, kind_color=None):
    ref(reference)
    label(kind, kind_color)
    rich(markup)


# ============================================================
# front matter
# ============================================================
para("HMIS Basic Navigation — Lesson 1: Finding a Participant",
     bold=True, size=20, color=TEAL, space_after=4)
para("Learner-facing text from the interactive tools", size=14, color=GREY, space_after=16)

h2("What this is")
para("Every word the learner reads inside the three simulations — Section 7, Section 10 and "
     "Section 11 — and the four step embeds that sit inside Section 11's slides. It is here to "
     "be rewritten in plain language and folded into the full Search Training script.")

h2("What is not in it")
para("The simulated interface's own wording: menu items, column headings, field labels, "
     "buttons, the empty state. That copy is Clarity's rather than ours, and the simulation is "
     "faithful to it on purpose — rewriting it would make the practice copy teach the wrong "
     "screen.")

h2("How to use it")
para("Each block carries a grey reference above it, e.g. S7.TASK.3.HINT. Edit the text under "
     "the reference and keep the reference intact; that is what an edit is applied against. "
     "The simulation source stays the source of truth, so edits made here are carried back "
     "into it rather than read out of this file.")
note("Generated from the built tools by script/extract_copy.mjs and script/make_copy_doc.py. "
     "Re-running either overwrites this file — work in a copy.")

h2("Who says what")
para("Lashes speaks the orientation, the scenario beats and the closings. The task blocks — "
     "situation, instruction, hint, feedback — are the training panel's own text rather than "
     "hers. Narration says participant; the simulated product says Client. Both are correct "
     "and the difference is deliberate.")

# ============================================================
# the sections
# ============================================================
SECTION_NOTE = {
    7:  "The learner's first contact with the simulation. The orientation runs before the first "
        "task and includes the only introduction Lashes gives.",
    10: "Verifying rather than finding: proving a record belongs to the person in front of you, "
        "and choosing when more than one matches.",
    11: "One participant, start to finish. The learner chooses which step to reach for, "
        "instead of being told.",
}

for sec in DATA["sections"]:
    n = sec["section"]
    h1({7: "Simulator 1 — Finding a Participant",
        10: "Simulator 2 — Verifying a Record"}.get(n, "Section %d" % n))
    para(sec["title"], color=GREY, space_after=4)
    if sec.get("scenarioTitle"):
        para("Scenario: %s" % sec["scenarioTitle"], bold=True, space_after=4)
    note(SECTION_NOTE.get(n, ""))

    if sec["orientation"]:
        h2("Orientation — spoken by Lashes")
        note("Plays once, before the first task. Beat 1 is the title card: she is drawn at full "
             "size with the words beside her and no bubble, so the section opens on a greeting.")
        for beat in sec["orientation"]:
            block("S%d.INTRO.%02d" % (n, beat["n"]),
                  "Beat %d%s" % (beat["n"],
                                 (" — button: %s" % beat["next"]) if beat["next"] else ""),
                  beat["html"])
        rule()
    else:
        h2("Orientation")
        para("None. This section opens straight onto the first task, with no introduction and "
             "no scene-setting.", color=RED)
        note("Flagged rather than silently omitted: if the full script gives Section 10 an "
             "opening, it has to be built, because there is nothing here to rewrite.")
        rule()

    if sec["scenarioOpen"]:
        h2("Scenario opening — spoken by Lashes")
        for beat in sec["scenarioOpen"]:
            block("S%d.OPEN.%02d" % (n, beat["n"]),
                  "Beat %d%s" % (beat["n"],
                                 (" — button: %s" % beat["next"]) if beat["next"] else ""),
                  beat["html"])
        rule()

    if sec["scenarioSteps"]:
        h2("Scenario step titles")
        note("Shown in the panel as the learner works through the list.")
        for i, t in enumerate(sec["scenarioSteps"], 1):
            ref("S%d.STEPTITLE.%d" % (n, i))
            rich("<p>%s</p>" % _html.escape(t))
        rule()

    h2("Tasks")
    note("Each task is four blocks: the situation the learner is put in, the instruction, the "
         "hint they can ask for at any time, and the feedback once they get it right.")
    for t in sec["tasks"]:
        para("Task %d — %s" % (t["n"], t["title"]), bold=True, size=12, space_after=4)
        if t["answer"] and t["answer"]["name"]:
            para("Correct record: %s (%s)" % (t["answer"]["name"], t["answer"]["id"]),
                 size=9, color=GREY, space_after=8)
        base = "S%d.TASK.%d" % (n, t["n"])
        block(base + ".SITUATION", "Situation", t["situation"])
        block(base + ".INSTRUCTION", "Instruction", t["instruction"])
        block(base + ".HINT", "Hint — on request", t["hint"])
        block(base + ".FEEDBACK", "Feedback — correct", t["feedback"], GREEN)
        for i, w in enumerate(t["wrong"], 1):
            block("%s.WRONG.%d" % (base, i),
                  "Feedback — opened %s" % (w["name"] or w["id"]), w["text"], ORANGE)
        rule()

    if sec["interstitial"]:
        h2("Interstitial — %s" % sec["interstitial"]["title"])
        note("Shown between the two halves of the lesson when the sections are played as one.")
        block("S%d.INTERSTITIAL" % n, "Screen text", sec["interstitial"]["html"])
        rule()

    if sec["scenarioClose"]:
        h2("Scenario close — spoken by Lashes")
        for beat in sec["scenarioClose"]:
            block("S%d.CLOSE.%02d" % (n, beat["n"]),
                  "Beat %d%s" % (beat["n"],
                                 (" — button: %s" % beat["next"]) if beat["next"] else ""),
                  beat["html"])
        rule()

    if sec["closing"]:
        h2("Completion screen")
        note("Shown when every task in the section is done.")
        block("S%d.DONE" % n, "Screen text", sec["closing"])
        rule()

# ============================================================
# the step embeds
# ============================================================
h1("The task embeds")
para("Three separate tools, one per slide. Each is the interface and nothing else — the slide "
     "carries the instruction — and each waits for the learner to do the thing the slide "
     "describes and then says what happened.", color=GREY, space_after=10)
note("These are not scored. The gate is the search itself: nothing advances until the learner "
     "has run it. An embed the script gives nothing to say has no text here, and says nothing.")

for st_ in DATA["steps"]:
    h2("Step %s  (%s)" % (st_["key"], st_["file"]))
    for ph in st_["phases"]:
        base = "STEP%s.%d" % (st_["key"].replace(".", "-"), ph["n"])
        block(base + ".HAPPENS", "What happens", ph["happens"])
    rule()

# ============================================================
# everything the learner can meet in any section
# ============================================================
h1("Shared text")
para("Written once and reused across all three sections.", color=GREY, space_after=12)

h2("Feedback bubbles")
note("The task's own feedback follows the Correct line above; these are the standing parts.")
for i, f in enumerate(DATA["feedback"], 1):
    colour = GREEN if f["kind"] == "Correct" else (ORANGE if f["kind"] == "Incorrect" else TEAL)
    ref("SHARED.FEEDBACK.%d" % i)
    label("%s — heading: %s" % (f["kind"], f["title"]), colour)
    rich("<p>%s</p>" % _html.escape(f["body"]))
rule()

h2("Completion screen — standing paragraph")
note("Sits under whichever section closing is shown.")
ref("SHARED.DONE.HABIT")
label("Heading: The habit underneath all of it")
rich("<p>%s</p>" % _html.escape(DATA["panel"]["habit"]))
rule()

h2("The training panel")
note("The panel is ours rather than Clarity's, so its wording is learner-facing copy too.")
ref("SHARED.PANEL.HELP")
label("Hover on the ? in the panel bar")
rich(DATA["panel"]["hover"])
ref("SHARED.PANEL.BUTTONS")
label("Buttons")
rich("<p>%s</p><p>%s</p>" % (_html.escape(DATA["panel"]["hint"]),
                             _html.escape(DATA["panel"]["next"])))

d.save(OUT)

counts = {
    "sections": len(DATA["sections"]),
    "tasks": sum(len(s["tasks"]) for s in DATA["sections"]),
    "beats": sum(len(s["orientation"]) + len(s["scenarioOpen"]) + len(s["scenarioClose"])
                 for s in DATA["sections"]),
    "steps": sum(len(s["phases"]) for s in DATA["steps"]),
}
print("  %s" % os.path.basename(OUT))
print("  %(sections)d sections, %(tasks)d tasks, %(beats)d spoken beats, "
      "%(steps)d step phases" % counts)
